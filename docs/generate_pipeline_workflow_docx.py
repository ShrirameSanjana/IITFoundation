#!/usr/bin/env python3
"""Generate PIPELINE_WORKFLOW.docx — pipeline workflow in WORKFLOW.docx format."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

DOCS_DIR = Path(__file__).resolve().parent
OUT_PATH = DOCS_DIR / "PIPELINE_WORKFLOW.docx"
OUT_PATH_FALLBACK = DOCS_DIR / "PIPELINE_WORKFLOW_new.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_flow_line(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Normal")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Foundation Textbook Pipeline", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Technical Workflow — What the Pipeline Does")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("June 2026")

    add_heading(doc, "Executive Summary", 1)
    add_para(
        doc,
        "This pipeline converts Foundation-series textbook PDFs into structured digital "
        "content: chapter summaries (student-friendly revision notes), key learning "
        "points, theory sections split by topic, and questions with answers grouped by "
        "type (illustrations, Check Your Knowledge, textbook exercises, practice "
        "exercises, case studies). Content is produced as JSON (format v3.1), exported "
        "to a flat QA table, optionally loaded into MySQL (database: foundation), and "
        "reviewed in local web viewers with PDF/HTML/JSON download.",
    )
    add_para(
        doc,
        "Critical point: topicwise_pipeline.py stops at JSON files. Database load "
        "(insert_qa_table.py) and viewing (viewer_api.py) are separate manual steps. "
        "Human review is recommended before treating any book as production-ready.",
        bold=True,
    )

    add_heading(doc, "End-to-End Workflow", 1)
    add_flow_line(doc, "PDF in Input_PDFs/")
    add_flow_line(doc, "│")
    add_flow_line(doc, "▼")
    add_flow_line(doc, "Mathpix OCR  ──►  Mathpix_Cache/<book>_mathpix.md")
    add_flow_line(doc, "│")
    add_flow_line(doc, "▼")
    add_flow_line(doc, "Split by chapter/topic  ──►  outputs/<book>/topics_md/")
    add_flow_line(doc, "│")
    add_flow_line(doc, "▼")
    add_flow_line(
        doc,
        "Ollama (qwen3:8b) — key points + student summaries  ──►  topics_json/",
    )
    add_flow_line(doc, "│")
    add_flow_line(doc, "▼")
    add_flow_line(
        doc,
        "Merge + relabel (v3.1) + MathML  ──►  outputs/<book>/<book>_final.json",
    )
    add_flow_line(doc, "│                      └── <book>_qa_table.json  (DB-ready)")
    add_flow_line(doc, "▼")
    add_flow_line(doc, "insert_qa_table.py  ──►  MySQL (foundation database)")
    add_flow_line(doc, "│")
    add_flow_line(doc, "▼")
    add_flow_line(
        doc,
        "viewer_api.py + viewers  ──►  Browse / Summary PDF / Full PDF in browser",
    )

    add_heading(doc, "Database Structure", 1)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["qa_chapter", "Chapter header: name, pages, summary, key points"],
            ["qa_theory_chapter", "Theory subsections linked to chapter_id"],
            ["qa_content_row", "Individual questions and answers"],
        ],
    )
    add_para(doc, "Schema: schema/qa_theory_and_rows.sql  |  Database name: foundation")

    add_heading(doc, "How Users View Content", 1)
    add_para(doc, "Start the API server:", bold=True)
    add_para(doc, "python viewer_api.py")
    add_para(doc, "Two viewers (both require the API server — do not open HTML via file://):", bold=True)
    add_bullets(
        doc,
        [
            "DB viewer: http://127.0.0.1:8765/Viewer/textbook_viewer.html — chapters from MySQL",
            "JSON viewer: http://127.0.0.1:8765/Viewer/output_json_viewer.html?json=/outputs/<book>/<book>_final.json",
        ],
    )
    add_bullets(
        doc,
        [
            "Book dropdown lists all book_slug values in qa_chapter",
            "Chapter list and tabs load data via /api/books, /api/chapters, /api/chapter/{id}",
            "Summary, Key points, Theory, Illustrations, Check knowledge, Textbook, Exercises tabs",
            "MathJax renders equations in theory and Q&A",
            "Download: Summary PDF, Full chapter PDF, HTML, JSON",
        ],
    )

    add_heading(doc, "Quick Reference", 1)
    add_para(doc, "1. Full pipeline from PDF:")
    add_para(doc, '   python -u topicwise_pipeline.py "Input_PDFs\\10 PHYSICS FOUNDATION.pdf" --with-images')
    add_para(doc, "2. Regenerate student summaries only:")
    add_para(
        doc,
        '   python -u topicwise_pipeline.py --summarize-only '
        '"outputs\\10 PHYSICS FOUNDATION\\10 PHYSICS FOUNDATION_final.json" '
        "--force-summarize --summarize-llm",
    )
    add_para(doc, "3. Load into MySQL:")
    add_para(
        doc,
        '   python insert_qa_table.py "outputs\\10 PHYSICS FOUNDATION\\10 PHYSICS FOUNDATION_qa_table.json"',
    )
    add_para(doc, "4. Start viewer API:")
    add_para(doc, "   python viewer_api.py")
    add_para(doc, "5. Open browser:")
    add_para(doc, "   http://127.0.0.1:8765/Viewer/textbook_viewer.html")
    add_para(
        doc,
        "   http://127.0.0.1:8765/Viewer/output_json_viewer.html"
        "?json=/outputs/10%20PHYSICS%20FOUNDATION/10%20PHYSICS%20FOUNDATION_final.json",
    )

    return doc


def main() -> None:
    doc = build_document()
    try:
        doc.save(OUT_PATH)
        print(f"Wrote {OUT_PATH}")
    except PermissionError:
        doc.save(OUT_PATH_FALLBACK)
        print(f"Wrote {OUT_PATH_FALLBACK} (close PIPELINE_WORKFLOW.docx in Word, then rename)")


if __name__ == "__main__":
    main()
