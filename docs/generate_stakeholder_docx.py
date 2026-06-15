#!/usr/bin/env python3
"""Generate STAKEHOLDER_WORKFLOW.docx from structured content."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

DOCS_DIR = Path(__file__).resolve().parent
OUT_PATH = DOCS_DIR / "STAKEHOLDER_WORKFLOW.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val)
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Foundation Textbook Digitization", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Stakeholder Overview — What Is Done and How It Works")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Audience: Executives, curriculum leads, program sponsors")
    doc.add_paragraph("May 2026")

    add_heading(doc, "Executive Summary", 1)
    add_para(
        doc,
        "This project converts Foundation-series textbook PDFs into structured digital "
        "content: chapter summaries, key learning points, theory sections, and "
        "questions with answers. Content is produced as JSON, loaded into MySQL "
        "(database: foundation), and reviewed in a local web viewer. "
        "Latest snapshot: 17,932 questions with solutions across all loaded books.",
    )
    add_para(
        doc,
        "Critical point: The automated pipeline stops at JSON files. Database load "
        "and viewing are separate steps with human review recommended between them.",
        bold=True,
    )

    add_heading(doc, "What Has Been Delivered", 1)
    add_table(
        doc,
        ["Capability", "Status"],
        [
            ["PDF to Markdown (Mathpix OCR)", "Operational"],
            ["Per-chapter extraction and LLM enrichment", "Operational"],
            ["Book JSON and QA table export", "Operational"],
            ["MySQL storage (3 tables)", "Operational"],
            ["Web chapter viewer", "Operational (local)"],
            ["Public cloud deployment", "Not in scope"],
        ],
    )

    add_heading(doc, "Books Loaded into MySQL", 2)
    add_table(
        doc,
        ["Book", "Chapters", "Theory sections", "Q&A rows"],
        [
            ["10 PHYSICS FOUNDATION", "8", "639", "2,944"],
            ["10TH BIOLOGY FOUNDATION", "8", "412", "2,368"],
            ["10TH CHEMISTRY FOUNDATION", "9", "698", "2,366"],
            ["10TH MATHS FOUNDATION", "19", "701", "1,012"],
            ["Foundation Biology Class 8th", "11", "141", "32"],
            ["Foundation Chemistry Class 8", "8", "472", "322"],
            ["Foundation Mathematics Class 6", "14", "727", "330"],
            ["Foundation Mathematics Class 7", "15", "616", "970"],
            ["Foundation Mathematics Class 8", "16", "844", "160"],
            ["Foundation Physics Class 8", "11", "738", "760"],
            ["Foundation Science Class 6", "16", "804", "3,168"],
            ["Foundation Science Class 7", "18", "946", "3,780"],
        ],
    )
    add_para(
        doc,
        "These counts show coverage in the database, not editorial approval. "
        "Each book should be spot-checked in the viewer before production use.",
    )

    add_heading(doc, "Questions With Solutions (Database)", 2)
    add_table(
        doc,
        ["Metric", "Count"],
        [
            ["Total questions with solutions (all books)", "17,932"],
            ["Total Q&A rows in database", "18,212"],
            ["Questions without a solution", "280"],
        ],
    )

    add_heading(doc, "Class-wise (with solutions)", 3)
    add_table(
        doc,
        ["Class", "Count"],
        [
            ["Class 6", "3,474"],
            ["Class 7", "4,734"],
            ["Class 8", "1,212"],
            ["Class 10", "8,512"],
            ["Combined total", "17,932"],
        ],
    )

    add_heading(doc, "Subject-wise (with solutions)", 3)
    add_table(
        doc,
        ["Subject", "Count"],
        [
            ["Science", "6,948"],
            ["Physics", "3,572"],
            ["Chemistry", "2,604"],
            ["Biology", "2,400"],
            ["Mathematics", "1,412"],
            ["Maths", "996"],
            ["Mathematics + Maths combined", "2,408"],
            ["Combined total", "17,932"],
        ],
    )

    add_heading(doc, "Per book (with solutions / total Q&A)", 3)
    add_table(
        doc,
        ["Book", "With solution", "Total Q&A"],
        [
            ["Foundation Science Class 7", "3,780", "3,780"],
            ["Foundation Science Class 6", "3,168", "3,168"],
            ["10 PHYSICS FOUNDATION", "2,866", "2,944"],
            ["10TH BIOLOGY FOUNDATION", "2,368", "2,368"],
            ["10TH CHEMISTRY FOUNDATION", "2,282", "2,366"],
            ["10TH MATHS FOUNDATION", "996", "1,012"],
            ["Foundation Mathematics Class 7", "954", "970"],
            ["Foundation Physics Class 8", "706", "760"],
            ["Foundation Chemistry Class 8", "322", "322"],
            ["Foundation Mathematics Class 6", "306", "330"],
            ["Foundation Mathematics Class 8", "152", "160"],
            ["Foundation Biology Class 8th", "32", "32"],
        ],
    )
    add_para(
        doc,
        "A row counts as having a solution when answer is not null and not blank. "
        "Snapshot: May 2026, database foundation, table qa_content_row.",
    )

    add_heading(doc, "End-to-End Workflow", 1)
    add_para(doc, "Phase A — Content pipeline (topicwise_pipeline.py), six steps:")
    add_bullets(
        doc,
        [
            "Step 1: Resolve source PDF (Input_PDFs/)",
            "Step 2: Mathpix OCR to cached Markdown (Mathpix_Cache/)",
            "Step 3: Split into per-topic Markdown (outputs/<book>/topics_md/)",
            "Step 4–5: Extract content; Ollama adds summary and key points (topics_json/)",
            "Step 6: Merge to <book>_final.json and <book>_qa_table.json",
        ],
    )
    add_para(doc, "Example command:")
    add_para(doc, 'python topicwise_pipeline.py "Input_PDFs\\10 PHYSICS FOUNDATION.pdf"')

    add_para(doc, "Phase B — Database load (insert_qa_table.py):")
    add_bullets(
        doc,
        [
            "Reads *_qa_table.json for one book",
            "Upserts qa_chapter and qa_theory_chapter",
            "Inserts Q&A rows into qa_content_row",
            "Use --replace-book to delete and reload a book cleanly",
        ],
    )
    add_para(doc, "Example command:")
    add_para(
        doc,
        'python insert_qa_table.py "outputs\\10 PHYSICS FOUNDATION\\10 PHYSICS FOUNDATION_qa_table.json"',
    )

    add_para(doc, "Phase C — Review (viewer_api.py + textbook_viewer.html):")
    add_bullets(
        doc,
        [
            "Start API: python viewer_api.py",
            "Open: http://127.0.0.1:8765/Viewer/textbook_viewer.html",
            "Select book, chapter; browse Summary, Key points, Theory, Q&A tabs",
        ],
    )

    add_heading(doc, "Database Structure", 1)
    add_table(
        doc,
        ["Table", "Purpose"],
        [
            ["qa_chapter", "Chapter header: name, pages, summary, key points"],
            ["qa_theory_chapter", "Theory subsections linked to chapter_id"],
            ["qa_content_row", "Individual questions and answers"],
        ],
    )
    add_para(doc, "Schema: schema/qa_theory_and_rows.sql  |  Database name: foundation")

    add_heading(doc, "How Data Is Inserted", 1)
    add_bullets(
        doc,
        [
            "Pipeline produces <book>_qa_table.json (not inserted automatically)",
            "insert_qa_table.py maps topics[] to qa_chapter and qa_theory_chapter",
            "rows[] become qa_content_row records",
            "Set DB_PASSWORD (or --password) if MySQL access is denied",
            "Re-run with --replace-book after pipeline updates to avoid duplicate Q&A",
        ],
    )

    add_heading(doc, "How Users View Content", 1)
    add_bullets(
        doc,
        [
            "Run python viewer_api.py from repository root",
            "Open browser URL above (do not use file:// on the HTML)",
            "Book dropdown lists all book_slug values in qa_chapter",
            "Chapter list and tabs load data via /api/books, /api/chapters, /api/chapter/{id}",
            "MathJax renders equations in theory and Q&A",
        ],
    )

    add_heading(doc, "Quality Gates (Recommended)", 1)
    add_bullets(
        doc,
        [
            "Spot-check 3 chapters per book in all four tabs",
            "Verify math and symbols render correctly",
            "Compare chapter titles and page ranges to printed TOC",
            "Sample-check exercise answers against source PDF",
            "Document approval per book before calling content production-ready",
        ],
    )

    add_heading(doc, "Risks and Limitations", 1)
    add_bullets(
        doc,
        [
            "OCR and LLM output may contain errors; human review required",
            "Pipeline re-runs do not update MySQL until insert_qa_table.py runs again",
            "Viewer is local-only (127.0.0.1), not a hosted production app",
            "Images skipped by default in pipeline for speed",
            "MySQL data can be overwritten with --replace-book",
        ],
    )

    add_heading(doc, "Council Verdict (Documentation Approach)", 1)
    add_para(
        doc,
        "Consensus: Stakeholders need outcomes and review gates, not only technical steps. "
        "Strongest dissent: Skeptic warns against treating database row counts as curriculum sign-off. "
        "Recommendation: Use the per-book status table and quality checklist before go-live.",
    )

    add_heading(doc, "Quick Reference", 1)
    add_para(doc, "1. python topicwise_pipeline.py \"Input_PDFs\\<book>.pdf\"")
    add_para(doc, "2. python insert_qa_table.py \"outputs\\<book>\\<book>_qa_table.json\"")
    add_para(doc, "3. python viewer_api.py")
    add_para(doc, "4. Browser: http://127.0.0.1:8765/Viewer/textbook_viewer.html")

    doc.add_paragraph()
    p = doc.add_paragraph("Companion file: STAKEHOLDER_WORKFLOW.md (full markdown version)")
    p.italic = True

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
